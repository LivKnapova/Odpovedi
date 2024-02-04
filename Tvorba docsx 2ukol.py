from docx import Document
from docx.shared import Cm

# Set arguments
name = str(input('Zadajte meno a příjmení: '))
rod_cislo = str(input('Zadajte rodné číslo: '))
datum_odberu = str(input('Napište datum obděru: '))

document = Document()

# Set font
style = document.styles['Normal']
style.font.name = 'Calibri'

# Add title
document.add_heading('\t\t\tVýsledný protokol genetického vyšetření', level=1)


records = (
    ('Jméno a příjmení:	', name),
    ('Rodné číslo: ', rod_cislo),
    ('Datum odběru: ', datum_odberu)
)

# Set table
table = document.add_table(rows=3, cols=2)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Jméno a příjmení:'
hdr_cells[1].text = name

hdr_cells = table.rows[1].cells
hdr_cells[0].text = 'Rodné číslo:'
hdr_cells[1].text = rod_cislo

hdr_cells = table.rows[2].cells
hdr_cells[0].text = 'Datum odběru:'
hdr_cells[1].text = datum_odberu 


document.add_page_break()

# Save result
document.save('Výsledný protokol.docx')